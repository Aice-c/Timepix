"""
生成中期检查报告 Word 文档
字体：中文宋体，英文/数字 Times New Roman
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = r"d:\Project\Timepix\document"
os.makedirs(OUTPUT_DIR, exist_ok=True)
DOC_PATH = os.path.join(OUTPUT_DIR, "中期检查报告.docx")


def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size=Pt(12), bold=False):
    """设置 run 的中英文字体。"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), cn_font)


def add_paragraph(doc, text, style=None, cn_font="宋体", en_font="Times New Roman",
                  size=Pt(12), bold=False, alignment=None, first_line_indent=None,
                  space_before=Pt(0), space_after=Pt(6)):
    """添加段落，自动设置字体。"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, cn_font, en_font, size, bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def add_heading_paragraph(doc, text, level=1):
    """添加标题段落（手动格式化，不使用内置 Heading 样式以确保字体一致）。"""
    if level == 0:
        p = add_paragraph(doc, text, size=Pt(18), bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=Pt(12), space_after=Pt(12))
    elif level == 1:
        p = add_paragraph(doc, text, size=Pt(15), bold=True,
                          space_before=Pt(18), space_after=Pt(6))
    elif level == 2:
        p = add_paragraph(doc, text, size=Pt(13), bold=True,
                          space_before=Pt(12), space_after=Pt(6))
    else:
        p = add_paragraph(doc, text, size=Pt(12), bold=True,
                          space_before=Pt(6), space_after=Pt(6))
    return p


def add_body(doc, text):
    """添加正文段落，首行缩进两字符。"""
    return add_paragraph(doc, text, size=Pt(12),
                         first_line_indent=Cm(0.74),
                         space_after=Pt(6))


def add_bullet(doc, text, indent_level=0):
    """添加列表项（用符号模拟）。"""
    prefix = "  " * indent_level
    markers = ["•", "–", "◦"]
    marker = markers[min(indent_level, len(markers) - 1)]
    return add_paragraph(doc, f"{prefix}{marker} {text}", size=Pt(12),
                         first_line_indent=Cm(0.74),
                         space_after=Pt(3))


def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=Pt(11), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(11))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 列宽
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = w

    doc.add_paragraph()  # 表后空行
    return table


# ═══════════════════════════════════════════════════
#  构建文档
# ═══════════════════════════════════════════════════

doc = Document()

# 全局默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# 行距 1.5 倍
from docx.oxml import OxmlElement
pPrDefault = doc.styles['Normal'].element
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:line'), '360')  # 1.5倍行距 = 360 twips
spacing.set(qn('w:lineRule'), 'auto')


# ─── 标题 ───
add_heading_paragraph(doc, "本科毕业论文中期检查报告", level=0)

# 课题名称
p = doc.add_paragraph()
run1 = p.add_run("课题名称：")
set_run_font(run1, size=Pt(12), bold=True)
run2 = p.add_run("基于深度学习的 Timepix3 探测器带电粒子入射角识别方法研究")
set_run_font(run2, size=Pt(12))
p.paragraph_format.space_after = Pt(12)


# ═══════════════════════════════════════════════════
#  一、课题主要任务
# ═══════════════════════════════════════════════════
add_heading_paragraph(doc, "一、课题主要任务", level=1)

add_body(doc, "本课题以 Timepix3 混合像素探测器为基础，研究利用深度学习方法对带电粒子"
         "（α 粒子、质子）入射极角进行自动识别的可行性与性能边界。"
         "Timepix3 探测器具备 256×256 像素阵列（像素间距 55 μm），"
         "能够同时记录每个像素的到达时间（ToA）和能量沉积（ToT）信息，"
         "从而获取粒子击中探测器时形成的微观径迹簇图像。")

add_body(doc, "课题的主要任务包括以下几个方面：")

add_body(doc, "1. 数据预处理与特征分析：对实验采集的多角度（10°–90°）带电粒子 ToT 数据"
         "进行清洗、裁剪与标准化处理，提取径迹簇的几何、能量分布及空间梯度等手工特征，"
         "并通过统计检验评估不同角度类别间的特征可区分性。")

add_body(doc, "2. 深度学习模型构建与优化：构建基于卷积神经网络的角度分类/回归模型，"
         "系统对比不同网络架构（ResNet-18、ShallowResNet 等）、不同损失函数"
         "（交叉熵、EMD 距离损失）及不同标签策略（one-hot 硬标签、高斯软标签）的组合方案，"
         "通过消融实验确定最优配置。")

add_body(doc, "3. 角度分辨极限分析：针对接近垂直入射（80°–90°）区间模型性能急剧下降的现象，"
         "从探测器物理机制角度分析其根本原因，定量评估 Timepix3 探测器在极角识别任务上的"
         "物理分辨极限。")

add_body(doc, "4. 论文撰写：在实验与分析工作基础上，完成本科毕业论文的撰写。")


# ═══════════════════════════════════════════════════
#  二、简述开题以来所做的具体工作和进展
# ═══════════════════════════════════════════════════
add_heading_paragraph(doc, "二、简述开题以来所做的具体工作和进展", level=1)

add_body(doc, "开题至今，课题工作主要围绕文献调研、数据分析、模型实验三个维度展开，"
         "取得了以下具体进展：")

# 2.1
add_heading_paragraph(doc, "2.1 文献调研与方法对比", level=2)

add_body(doc, "精读了课题组推荐的核心参考文献（Okabe et al., Nature Communications, 2024），"
         "该文献利用仅 4 个 CZT 晶体像素构成的 Tetris 形状探测器，"
         "结合 Filter Layer + 1D U-Net 架构与 Wasserstein 距离损失函数，"
         "实现了 γ 射线方位角约 1° 的预测精度。")

add_body(doc, "在此基础上，系统分析了文献方法与本课题任务之间的核心差异：")

add_table(doc,
    headers=["对比维度", "参考文献", "本课题"],
    rows=[
        ["探测对象", "γ 射线（0.5 MeV）", "带电粒子（α / 质子）"],
        ["方向信息来源", "宏观几何遮挡\n（像素间铅屏蔽）", "微观径迹形态\n（硅层能量沉积）"],
        ["角度类型", "方位角 θ ∈ [0°, 360°)", "极角 θ ∈ [10°, 90°]"],
        ["输入维度", "4–6 个像素计数值", "50×50 矩阵\n（~12 个活跃像素）"],
        ["数据来源", "蒙特卡罗模拟", "真实实验数据"],
    ],
    col_widths=[Cm(3), Cm(5), Cm(5)])

add_body(doc, "该对比分析明确了直接移植文献方法（如 U-Net 架构、密集角度分 bin）的局限性，"
         "同时提炼出三项可迁移的核心思想加以借鉴。")

# 2.2
add_heading_paragraph(doc, "2.2 损失函数与标签策略改进", level=2)

add_body(doc, "针对传统交叉熵损失在有序角度分类任务中的不足"
         "（所有误分类等价惩罚，无法编码角度有序性），借鉴参考文献引入两项改进：")

add_bullet(doc, "EMD 损失函数（Earth Mover\u2019s Distance）：通过累积分布函数差值度量"
           "预测分布与目标分布之间的距离，使\u201c偏差 2\u00b0\u201d的损失小于\u201c偏差 10\u00b0\u201d，"
           "为模型提供与物理意义一致的梯度方向。")

add_bullet(doc, "高斯软标签：将 one-hot 硬标签替换为以真实角度为中心的高斯分布，"
           "缓解相邻角度数据高度相似时产生的梯度矛盾问题。")

# 2.3
add_heading_paragraph(doc, "2.3 网络架构适配", level=2)

add_body(doc, "分析发现标准 ResNet-18（为 224×224 密集图像设计，含 5 次下采样）"
         "不适合本课题的稀疏小尺寸输入（50×50，仅约 12 个活跃像素）："
         "过多的下采样操作导致有效空间信息在网络中段即被完全压缩。"
         "为此，设计了 ShallowResNet 浅层残差网络，仅保留 1 次下采样，"
         "以 3–4 层卷积覆盖径迹簇尺度的感受野，减少信息丢失。")

# 2.4
add_heading_paragraph(doc, "2.4 近垂直角度（80°–90°）数据可区分性分析", level=2)

add_body(doc, "对 80°–90°（2° 间隔，6 个类别）约 19.1 万条质子 ToT 数据"
         "开展了两阶段系统分析：")

add_body(doc, "第一阶段：提取 21 个基础手工特征（几何特征、ToT 统计量、能量分布特征），"
         "进行两两 KS 检验，所有相邻角度对的最大 KS 统计量均小于 0.035"
         "（显著性阈值 0.10），差异不显著；随机森林分类器准确率仅 19.0%"
         "（随机基线 16.7%）。")

add_body(doc, "第二阶段：补充 32 个高阶空间特征（梯度特征、方位角特征、Hu 矩等），"
         "总特征数达 53 个。最大 KS 统计量仅微升至 0.034，随机森林准确率 20.0%，"
         "改善有限。")

add_body(doc, "以上分析从统计学层面量化证实：在 2° 间隔条件下，"
         "80°–90° 各角度类别的数据分布在已知特征空间中高度重叠，缺乏有效判别信息。")

# 2.5
add_heading_paragraph(doc, "2.5 全范围角度消融实验", level=2)

add_body(doc, "在 9 个角度类别（10°, 20°, 30°, 45°, 50°, 60°, 70°, 80°, 90°）上，"
         "完成了 2×3 消融对比实验：")

add_table(doc,
    headers=["配置", "模型", "损失函数/策略", "准确率", "Macro-F1", "MAE(°)"],
    rows=[
        ["A", "ResNet-18", "交叉熵 + One-hot", "66.6%", "—", "—"],
        ["B", "ResNet-18", "EMD + 高斯软标签", "67.84%", "提升", "4.45"],
        ["C", "ShallowResNet", "交叉熵 + One-hot", "65.5%", "—", "—"],
        ["D", "ShallowResNet", "EMD + 高斯软标签", "66.2%", "提升", "—"],
        ["E", "ResNet-18", "MSE 回归", "≈65%等效", "—", "—"],
        ["F", "ShallowResNet", "MSE 回归", "≈60%等效", "—", "—"],
    ],
    col_widths=[Cm(1.2), Cm(2.8), Cm(3.5), Cm(2.0), Cm(1.8), Cm(1.8)])

add_body(doc, "主要发现：")

add_body(doc, "（1）EMD 损失函数在两种架构上均带来正向提升，"
         "其中在 ShallowResNet 上提升更为明显（+0.7%），在 ResNet-18 上亦有效。")

add_body(doc, "（2）ResNet-18 仍优于 ShallowResNet 约 1%，"
         "推测原因在于 ResNet-18 参数量更大（~11M vs ~1M），能学到更多有效特征。")

add_body(doc, "（3）分类方案整体优于回归方案，"
         "最佳分类准确率 67.84% 高于最佳回归等效准确率 65%。")

add_body(doc, "（4）误差主要集中在 70°–90° 区间，"
         "10°–70° 各类别分类准确率良好，80°–90° 之间存在严重混淆。")


# ═══════════════════════════════════════════════════
#  三、存在的主要问题
# ═══════════════════════════════════════════════════
add_heading_paragraph(doc, "三、存在的主要问题", level=1)

# 3.1
add_heading_paragraph(doc, "3.1 近垂直角度分类精度受限于探测器物理分辨率", level=2)

add_body(doc, "全范围实验的最佳准确率为 67.84%，整体精度受限的核心原因在于模型在 80°–90° "
         "区间的混淆。经分析，该限制并非模型能力不足，而是源于探测器固有的物理约束：")

add_bullet(doc, "投影长度差异极小：80° 和 90° 入射粒子在 300 μm 厚硅传感器中的"
           "径迹投影长度差异不足 1 个像素宽度（55 μm）。")

add_bullet(doc, "电荷扩散效应掩盖角度信号：电荷横向扩散尺度约 3–4 个像素"
           "（165–220 μm），远大于相邻角度导致的投影差异量级"
           "（约 4%–5% 的特征变化幅度），使得角度信息被统计涨落淹没。")

add_body(doc, "该物理限制已通过多种独立方法交叉验证"
         "（KS 检验、随机森林、深度学习模型均一致表明 80°–90° 不可区分），"
         "表征的是探测器的角分辨极限，而非算法的局限。")

# 3.2
add_heading_paragraph(doc, "3.2 实验数据的固有约束", level=2)

add_body(doc, "与参考文献使用蒙特卡罗模拟数据（可无限生成、精确控制角度）不同，"
         "本课题依赖真实实验数据，存在以下约束：样本量有限、角度采样不均匀且离散"
         "（仅 9 个固定角度）、包含实验噪声与统计涨落。"
         "这些因素共同限制了部分高级方法（如密集角度分 bin 回归）的适用性。")

# 3.3
add_heading_paragraph(doc, "3.3 学术论文图表规范性不足", level=2)

add_body(doc, "当前实验结果的可视化图表在子图编号布局、图例文字大小、"
         "数据标签重叠等方面存在不规范之处，需参照高水平期刊论文标准进行优化。")


# ═══════════════════════════════════════════════════
#  四、下一步主要研究任务和安排
# ═══════════════════════════════════════════════════
add_heading_paragraph(doc, "四、下一步主要研究任务和安排", level=1)

# 4.1
add_heading_paragraph(doc, "4.1 完善 80°–90° 不可区分的物理解释（4 月中旬）", level=2)

add_body(doc, "进一步调研近垂直角度粒子在硅探测器中的能量沉积物理机制，"
         "结合 Timepix3 探测器的像素尺寸、传感器厚度、电荷扩散模型等参数，"
         "完善角度分辨极限的定量解释，使结论具有物理学上的严谨性。")

# 4.2
add_heading_paragraph(doc, "4.2 优化实验图表规范性（4 月中旬）", level=2)

add_body(doc, "参照高水平学术期刊的图表标准，重新绘制混淆矩阵、"
         "分类指标对比图等核心可视化内容，规范子图标号（A, B, C, D）、"
         "图例大小与注释布局。")

# 4.3
add_heading_paragraph(doc, "4.3 撰写论文初稿（4 月中旬至 5 月初）", level=2)

add_body(doc, "按照以下结构撰写中文版毕业论文初稿：")

add_body(doc, "（1）引言：研究背景、文献综述、课题意义与研究目标。")
add_body(doc, "（2）实验方法：探测器简介、数据集说明、数据预处理流程、"
         "网络架构设计（ResNet-18、ShallowResNet）、损失函数与标签策略"
         "（交叉熵 / EMD / 高斯软标签）、消融实验设计。")
add_body(doc, "（3）实验结果与分析：分层展示——先呈现 10°–70° 区间的良好分类表现，"
         "再详细分析 80°–90° 区间准确率下降的物理原因，"
         "最后呈现消融实验各配置对比结果。")
add_body(doc, "（4）讨论：角度分辨极限的物理意义、方法改进的有效性评估、"
         "与参考文献方法的差异分析、前期探索过的其他方法"
         "（如小样本学习方案）的尝试与分析。")
add_body(doc, "（5）结论与展望。")

# 4.4
add_heading_paragraph(doc, "4.4 修改与完善（5 月）", level=2)

add_body(doc, "根据导师反馈意见对论文初稿进行修改、完善结果展示与讨论深度，"
         "准备答辩材料。")


# ═══════════════════════════════════════════════════
#  保存
# ═══════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f"Document saved to: {DOC_PATH}")
